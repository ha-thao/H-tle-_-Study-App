# study_app.py - Hoàn chỉnh (đã sửa lỗi folder không đều + góc nhọn)
import customtkinter as ctk
from tkinter import messagebox, filedialog, PanedWindow
from PIL import Image
import os
import time
import shutil
from datetime import datetime

try:
    import fitz
    import docx
    HAS_LIBS = True
except ImportError:
    HAS_LIBS = False
    print("CẢNH BÁO: Thiếu PyMuPDF hoặc python-docx. Cài: pip install PyMuPDF python-docx")

ctk.set_appearance_mode("Light")
ctk.set_default_color_theme("blue")


class StudyAppPro(ctk.CTkFrame):
    def __init__(self, parent, db, mssv, ho_ten, theme):
        super().__init__(parent, fg_color="transparent")
        self.parent = parent
        self.db = db
        self.mssv = mssv
        self.ho_ten = ho_ten
        self.theme = theme

        self.bg_color = self.c("#FDF6E3")
        self.accent_color = self.c("#3F51B5")
        self.card_color = self.c("#FFFFFF")
        self.folder_color = self.c("#D1D5F0")
        
        self.current_semester = 1
        self.current_semester_name = "Học kỳ 1"
        
        self.semester_var = ctk.StringVar(value=self.current_semester_name)

        self.semesters = []
        self.subjects_data = {}
        self.files_data = {}
        self.notes_data = {}

        self.current_semester = None
        self.current_subject = None
        self.current_tab = "files"
        self.current_file = None
        self.current_note = None
        self.view_state = "HOME"
        self.pdf_zoom = 1.2
        self.search_keyword = ""
        self.search_file_keyword = ""
        self.user_set_split_pos = None

        self.upload_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "uploads")
        if not os.path.exists(self.upload_dir):
            os.makedirs(self.upload_dir)

        self.load_data_from_db()
        self.load_assets()
        self.setup_ui()

    def c(self, hex_code):
        if self.theme:
            return self.theme.get(hex_code, hex_code)
        return hex_code

    # ==================== LOAD DATA ====================
    def load_data_from_db(self):
        try:
            # Tạo cứng 7 học kỳ
            self.semesters = [{"id_hoc_ky": i, "hoc_ky": i, "ten_hoc_ky": f"Học kỳ {i}"} for i in range(1, 9)]
            self.current_semester = 1
            self.current_semester_name = "Học kỳ 1"
            
            # Reset data
            self.subjects_data = {}
            self.files_data = {}
            self.notes_data = {}
            
            for sem in self.semesters:
                ky = sem['hoc_ky']
                courses = self.db.get_courses_by_semester(ky)
                self.subjects_data[ky] = []
                for course in courses:
                    subject_name = course[1]
                    self.subjects_data[ky].append({
                        "id": course[0],
                        "name": subject_name,
                        "tin_chi": course[2],
                        "loai": course[3],
                        "da_hoc": course[4],
                        "diem_tb": course[5]
                    })
                    
                    self.files_data[subject_name] = []
                    self.notes_data[subject_name] = []
                    
                    try:
                        files = self.db.get_files(subject_name)
                        for f in files:
                            file_path = f[4] if len(f) > 4 else ""
                            size_str = "0KB"
                            if file_path and os.path.exists(file_path):
                                sz = os.path.getsize(file_path) // 1024
                                size_str = f"{sz}KB" if sz < 1024 else f"{sz/1024:.1f}MB"
                            
                            self.files_data[subject_name].append({
                                "id": f[0],
                                "name": f[1],
                                "type": f[2],
                                "date": f[3].strftime("%d/%m/%Y") if hasattr(f[3], 'strftime') else str(f[3]),
                                "size": size_str,
                                "path": file_path
                            })
                    except:
                        pass
                    
                    try:
                        notes = self.db.get_notes(subject_name)
                        for n in notes:
                            self.notes_data[subject_name].append({
                                "_id": n[0],
                                "title": n[1],
                                "content": n[2],
                                "time": n[3].strftime("%d/%m/%Y") if hasattr(n[3], 'strftime') else str(n[3])
                            })
                    except:
                        pass
            
            if self.semesters:
                self.current_semester = self.semesters[0]['hoc_ky']
                self.current_semester_name = self.semesters[0]['ten_hoc_ky']
        except Exception as e:
            print(f"Lỗi load data: {e}")
            self.semesters = [{"id_hoc_ky": 1, "hoc_ky": 1, "ten_hoc_ky": "Học kỳ 1"}]
            self.current_semester = 1

    def load_assets(self):
        self.folder_icon = None

    def create_folder_item(self, parent, name, r, c):
        folder = ctk.CTkFrame(parent, fg_color=self.folder_color, height=190, 
                              corner_radius=25, cursor="hand2")
        folder.grid(row=r, column=c, padx=15, pady=15, sticky="nsew")
        folder.grid_propagate(False)
        
        # Nút xóa - LUÔN HIỂN THỊ
        x_btn = ctk.CTkButton(folder, text="✕", width=25, height=25, fg_color="#E74C3C", 
                              hover_color="#C0392B", corner_radius=15, text_color="white",
                              font=("Arial", 10, "bold"),
                              command=lambda n=name: self.confirm_delete_subject(n))
        x_btn.place(relx=1.0, x=-10, y=10, anchor="ne")

        icon_lbl = ctk.CTkLabel(folder, text="📁", text_color=self.c("#2E3A6E"), font=("Garet Variable", 60))
        icon_lbl.pack(pady=(35, 10))

        name_lbl = ctk.CTkLabel(folder, text=name, font=("Garet Variable", 14, "bold"), 
                                text_color=self.c("#2E3A6E"), wraplength=180)
        name_lbl.pack(pady=5, padx=10)

        def on_enter(e):
            folder.configure(fg_color=self.accent_color)
            icon_lbl.configure(text_color="white")
            name_lbl.configure(text_color="white")

        def on_leave(e):
            folder.configure(fg_color=self.folder_color)
            icon_lbl.configure(text_color=self.c("#2E3A6E"))
            name_lbl.configure(text_color=self.c("#2E3A6E"))

        # Chỉ gắn hover cho folder, icon, name (KHÔNG cho x_btn)
        folder.bind("<Enter>", on_enter)
        folder.bind("<Leave>", on_leave)
        icon_lbl.bind("<Enter>", on_enter)
        icon_lbl.bind("<Leave>", on_leave)
        name_lbl.bind("<Enter>", on_enter)
        name_lbl.bind("<Leave>", on_leave)
        
        # Mở môn khi click
        folder.bind("<Button-1>", lambda e, n=name: self.open_subject(n))
        icon_lbl.bind("<Button-1>", lambda e, n=name: self.open_subject(n))
        name_lbl.bind("<Button-1>", lambda e, n=name: self.open_subject(n))
        
        # Ngăn click vào nút x không lan ra folder
        x_btn.bind("<Button-1>", lambda e: "break")
    # ==================== UI SETUP ====================
    def setup_ui(self):
        self.main_container = ctk.CTkFrame(self, fg_color="transparent")
        self.main_container.pack(fill="both", expand=True)
        self.main_view = ctk.CTkFrame(self.main_container, fg_color="transparent")
        self.main_view.pack(fill="both", expand=True)
        self.render_view()

    def render_view(self):
        for w in self.main_view.winfo_children():
            w.destroy()

        if self.view_state == "HOME":
            self.draw_home_view()
        elif self.view_state == "SUBJECT":
            self.draw_subject_view()
        elif self.view_state == "VIEWER":
            self.draw_inline_file_viewer()
        elif self.view_state == "NOTE_EDITOR":
            self.draw_inline_note_editor()
        elif self.view_state == "SPLIT":
            self.draw_inline_split_session()
            self.after(50, self.adjust_split_50_50)

    def set_view_state(self, state):
        self.view_state = state
        self.render_view()

    def go_to_home(self):
        self.current_subject = None
        self.current_file = None
        self.current_note = None
        self.set_view_state("HOME")

    def change_semester(self, choice):
        ky = int(choice.split(" ")[-1])
        self.current_semester = ky
        self.current_semester_name = choice
        self.render_view()

    def open_subject(self, subject_name):
        self.current_subject = subject_name
        self.current_tab = "files"
        self.set_view_state("SUBJECT")

    def switch_tab(self, mode):
        self.current_tab = mode
        self.render_view()

    # ==================== HOME VIEW ====================
    def draw_home_view(self):
        header = ctk.CTkFrame(self.main_view, fg_color="transparent")
        header.pack(fill="x", padx=40, pady=(40, 10))
        
        left_h = ctk.CTkFrame(header, fg_color="transparent")
        left_h.pack(side="left")
        ctk.CTkLabel(left_h, text="DANH SÁCH MÔN HỌC", font=("Garet Variable", 28, "bold"), 
                     text_color=self.accent_color).pack(anchor="w")
        
        sem_values = [f"Học kỳ {i}" for i in range(1, 9)]
        sem_menu = ctk.CTkOptionMenu(left_h, 
                                     values=sem_values, 
                                     variable=self.semester_var,
                                     command=self.change_semester, 
                                     fg_color=self.c("#F0F0F0"), 
                                     button_color=self.accent_color,
                                     button_hover_color=self.c("#2E3A6E"),
                                     dropdown_hover_color=self.c("#2E3A6E"),
                                     text_color="black",
                                     width=200,
                                     font=("Garet Variable", 13),
                                     dropdown_font=("Garet Variable", 13))
        sem_menu.pack(anchor="w", pady=(10, 0))
        
        search_frame = ctk.CTkFrame(header, fg_color="transparent")
        search_frame.pack(side="right")
        self.search_entry = ctk.CTkEntry(search_frame, placeholder_text="🔍 Tìm kiếm môn học...", 
                                         font=("Garet Variable", 13), width=250, height=35, border_color=self.c("#B0BCE6"))
        self.search_entry.pack(side="left", padx=10)
        self.search_entry.bind("<Return>", lambda e: self.search_subjects())
        ctk.CTkButton(search_frame, text="Tìm", width=60, height=35, fg_color=self.accent_color, hover_color=self.c("#2E3A6E"),
                      font=("Garet Variable", 13), command=self.search_subjects).pack(side="left")

        scroll = ctk.CTkScrollableFrame(self.main_view, fg_color="transparent")
        scroll.pack(fill="both", expand=True, padx=40, pady=10)
        for i in range(4):
            scroll.grid_columnconfigure(i, weight=1, uniform="col")
        
        # Cấu hình 4 cột, mỗi cột giãn đều
        for i in range(4):
            scroll.grid_columnconfigure(i, weight=1)

        subjects = self.subjects_data.get(self.current_semester, [])
        if self.search_keyword:
            subjects = [s for s in subjects if self.search_keyword.lower() in s['name'].lower()]

        r, c = 0, 0
        for sub in subjects:
            self.create_folder_item(scroll, sub['name'], r, c)
            c += 1
            if c > 3:
                c = 0
                r += 1
        
        self.create_add_subject_card(scroll, r, c)

    def search_subjects(self):
        self.search_keyword = self.search_entry.get().strip()
        self.render_view()

    def create_add_subject_card(self, parent, r, c):
        folder = ctk.CTkFrame(parent, fg_color="transparent", border_width=2, border_color=self.accent_color,
                              height=190, corner_radius=25, cursor="hand2")
        folder.grid(row=r, column=c, padx=15, pady=15, sticky="nsew")
        folder.grid_propagate(False)

        icon_lbl = ctk.CTkLabel(folder, text="+", font=("Garet Variable", 60), text_color=self.accent_color)
        icon_lbl.pack(pady=(25, 0))
        name_lbl = ctk.CTkLabel(folder, text="Thêm môn học", font=("Garet Variable", 14, "bold"), 
                                text_color=self.accent_color)
        name_lbl.pack(pady=10)

        def on_enter(e):
            folder.configure(fg_color=self.accent_color)
            icon_lbl.configure(text_color="white")
            name_lbl.configure(text_color="white")

        def on_leave(e):
            folder.configure(fg_color="transparent")
            icon_lbl.configure(text_color=self.accent_color)
            name_lbl.configure(text_color=self.accent_color)

        for w in [folder, icon_lbl, name_lbl]:
            w.bind("<Enter>", on_enter)
            w.bind("<Leave>", on_leave)
            w.bind("<Button-1>", lambda e: self.add_subject())

    def add_subject(self):
        dialog = ctk.CTkInputDialog(text="Nhập tên môn học mới:", title="Thêm môn học", font=("Garet Variable", 13))
        subject_name = dialog.get_input()
        if subject_name:
            subject_name = subject_name.strip()
            for sub in self.subjects_data.get(self.current_semester, []):
                if sub['name'].lower() == subject_name.lower():
                    messagebox.showwarning("Cảnh báo", "Môn học này đã tồn tại!")
                    return
            try:
                self.db.add_subject(subject_name, self.current_semester)
                self.load_data_from_db()
                messagebox.showinfo("Thành công", f"Đã thêm môn học: {subject_name}")
                self.render_view()
            except Exception as e:
                messagebox.showerror("Lỗi", f"Thêm môn thất bại: {e}")
    
    def confirm_delete_subject(self, subject_name):
        if messagebox.askyesno("Xác nhận", f"Xóa môn '{subject_name}' khỏi danh sách học tập?"):
            if self.db.delete_student_subject_by_name(subject_name, self.current_semester):
                self.load_data_from_db()
                self.render_view()

    # ==================== SUBJECT VIEW ====================
    def draw_subject_view(self):
        header = ctk.CTkFrame(self.main_view, fg_color="transparent")
        header.pack(fill="x", padx=40, pady=(40, 10))

        ctk.CTkButton(header, text="← Quay lại", width=100, fg_color=self.c("#3F51B5"), 
                      text_color="black", font=("Garet Variable", 13), hover_color=self.c("#D1D5F0"), command=self.go_to_home).pack(side="left")
        ctk.CTkLabel(header, text=self.current_subject, font=("Garet Variable", 24, "bold"), 
                     text_color=self.accent_color).pack(side="left", padx=25)

        tab_box = ctk.CTkFrame(header, fg_color="transparent")
        tab_box.pack(side="left", padx=20)
        for t_text, t_mode in [("📄 Tài liệu", "files"), ("📝 Ghi chú", "notes")]:
            active = self.current_tab == t_mode
            ctk.CTkButton(tab_box, text=t_text, width=120, fg_color="transparent",
                          font=("Garet Variable", 16, "bold" if active else "normal"),
                          hover_color=self.c("#D1D5F0"),
                          text_color=self.accent_color if active else "gray",
                          command=lambda m=t_mode: self.switch_tab(m)).pack(side="left", padx=10)

        if self.current_tab == "files":
            ctk.CTkButton(header, text="📥 Upload tài liệu", fg_color=self.accent_color, 
                          corner_radius=20, width=130, font=("Garet Variable", 12, "bold"), height=38,
                          hover_color=self.c("#D1D5F0"),
                          command=self.upload_file).pack(side="right", padx=10)
        else:
            ctk.CTkButton(header, text="+ Thêm", fg_color=self.accent_color, hover_color=self.c("#2E3A6E"), 
                          corner_radius=20, width=100, font=("Garet Variable", 12, "bold"), height=38,
                          command=self.show_add_options).pack(side="right", padx=10)

        search_frame = ctk.CTkFrame(header, fg_color="transparent")
        search_frame.pack(side="right", padx=20)
        self.file_search_entry = ctk.CTkEntry(search_frame, placeholder_text="🔍 Tìm kiếm...", 
                                              font=("Garet Variable", 13), width=180, height=35)
        self.file_search_entry.pack(side="left")
        self.file_search_entry.bind("<Return>", lambda e: self.search_items())
        ctk.CTkButton(search_frame, text="Tìm", width=60, height=35, fg_color=self.accent_color, hover_color=self.c("#2E3A6E"),
                      font=("Garet Variable", 13), command=self.search_items).pack(side="left", padx=5)

        scroll = ctk.CTkScrollableFrame(self.main_view, fg_color="transparent")
        scroll.pack(fill="both", expand=True, padx=40, pady=10)

        if self.current_tab == "files":
            items = self.files_data.get(self.current_subject, [])
            if self.search_file_keyword:
                items = [f for f in items if self.search_file_keyword.lower() in f['name'].lower()]
            for f in items:
                self.create_file_card(scroll, f)
        else:
            items = self.notes_data.get(self.current_subject, [])
            if self.search_file_keyword:
                items = [n for n in items if self.search_file_keyword.lower() in n['title'].lower()]
            for n in items:
                self.create_note_card(scroll, n)

    def show_add_options(self):
        dialog = ctk.CTkToplevel(self)
        dialog.title("Thêm mới")
        dialog.geometry("350x280")
        dialog.attributes("-topmost", True)
        dialog.transient(self)
        dialog.grab_set()
        dialog.focus_force()
        
        dialog.update_idletasks()
        x = self.winfo_toplevel().winfo_x() + (self.winfo_toplevel().winfo_width() // 2) - 175
        y = self.winfo_toplevel().winfo_y() + (self.winfo_toplevel().winfo_height() // 2) - 140
        dialog.geometry(f"+{x}+{y}")
        
        container = ctk.CTkFrame(dialog, fg_color=self.c("#FFFFFF"), corner_radius=20,
                                 border_width=2, border_color=self.accent_color)
        container.pack(fill="both", expand=True, padx=15, pady=15)
        
        header = ctk.CTkFrame(container, fg_color=self.accent_color, height=50, corner_radius=0)
        header.pack(fill="x")
        ctk.CTkLabel(header, text="CHỌN LOẠI", font=("Garet Variable", 16, "bold"), 
                     text_color="white").pack(pady=12)
        
        form = ctk.CTkFrame(container, fg_color="transparent")
        form.pack(fill="both", expand=True, padx=20, pady=20)
        
        ctk.CTkLabel(form, text="Bạn muốn tạo ghi chú bằng cách nào?", 
                     font=("Garet Variable", 13)).pack(pady=10)
        
        ctk.CTkButton(form, text="📝 Tạo văn bản ngay", text_color="white", fg_color=self.accent_color, hover_color=self.c("#2E3A6E"),
                      font=("Garet Variable", 13), height=45, corner_radius=15,
                      command=lambda: [dialog.destroy(), self.create_new_note()]).pack(fill="x", pady=8)
        
        ctk.CTkButton(form, text="📤 Upload file (txt, docx)", text_color="white", fg_color=self.accent_color, hover_color=self.c("#2E3A6E"),
                      font=("Garet Variable", 13), height=45, corner_radius=15,
                      command=lambda: [dialog.destroy(), self.upload_file_to_note()]).pack(fill="x", pady=8)
        
        cancel_btn = ctk.CTkButton(form, text="Hủy", fg_color="gray", 
                                    font=("Garet Variable", 13), height=40, corner_radius=15,
                                    command=dialog.destroy)
        cancel_btn.pack(fill="x", pady=15)

    def search_items(self):
        self.search_file_keyword = self.file_search_entry.get().strip()
        self.render_view()

    def create_file_card(self, parent, f):
        card = ctk.CTkFrame(parent, fg_color=self.c("#FFFFFF"), height=85, corner_radius=20, cursor="hand2")
        card.pack(fill="x", pady=8)
        card.pack_propagate(False)
        
        ctk.CTkLabel(card, text="📕" if f.get('type') == "pdf" else "📘", 
                     font=("Garet Variable", 28), width=65).pack(side="left", padx=15)
        ctk.CTkLabel(card, text=f['name'], font=("Garet Variable", 15, "bold"), 
                     text_color="#34495E").pack(side="left")
        ctk.CTkLabel(card, text=f.get('size', '0KB'), font=("Garet Variable", 12), 
                     text_color="gray").pack(side="right", padx=30)
        ctk.CTkLabel(card, text=f.get('date', ''), font=("Garet Variable", 11), 
                     text_color="gray").pack(side="right", padx=10)

        btn_frame = ctk.CTkFrame(card, fg_color="transparent")
        btn_frame.pack(side="right", padx=10)
        
        ctk.CTkButton(btn_frame, text="Xem", width=50, font=("Garet Variable", 13),fg_color=self.c("#3F51B5"),  hover_color=self.c("#D1D5F0"),
                      command=lambda: self.switch_to_viewer(f)).pack(side="left", padx=2)
        ctk.CTkButton(btn_frame, text="Xóa", width=50, fg_color="red", font=("Garet Variable", 13), hover_color=self.c("#D1D5F0"),
                      command=lambda file=f: self.delete_file(file)).pack(side="left", padx=2)
        
        card.bind("<Button-1>", lambda e: self.switch_to_viewer(f))

    def delete_file(self, file_info):
        if messagebox.askyesno("Xác nhận", f"Xóa tài liệu '{file_info['name']}'?"):
            try:
                if file_info.get('path') and os.path.exists(file_info['path']):
                    os.remove(file_info['path'])
                self.db.delete_file(file_info['id'])
                if self.current_subject in self.files_data:
                    self.files_data[self.current_subject] = [f for f in self.files_data[self.current_subject] 
                                                              if f['id'] != file_info['id']]
                messagebox.showinfo("Thành công", "Đã xóa tài liệu!")
                self.render_view()
            except Exception as e:
                messagebox.showerror("Lỗi", f"Xóa thất bại: {e}")

    def create_note_card(self, parent, n):
        card = ctk.CTkFrame(parent, fg_color=self.c("#FFFFFF"), height=110, corner_radius=20, cursor="hand2")
        card.pack(fill="x", pady=10)
        card.pack_propagate(False)
        
        ctk.CTkLabel(card, text="📝", font=("Garet Variable", 30), width=65).pack(side="left", padx=15)
        info = ctk.CTkFrame(card, fg_color="transparent")
        info.pack(side="left", fill="both", expand=True, pady=18)
        ctk.CTkLabel(info, text=n['title'], font=("Garet Variable", 16, "bold"), 
                     text_color="#2C3E50").pack(anchor="w")
        ctk.CTkLabel(info, text=n['content'][:75] + "..." if len(n['content']) > 75 else n['content'], 
                     font=("Garet Variable", 13), text_color="gray").pack(anchor="w")
        
        btn_frame = ctk.CTkFrame(card, fg_color="transparent")
        btn_frame.pack(side="right", padx=10)
        ctk.CTkButton(btn_frame, text="Sửa", width=50, font=("Garet Variable", 13),fg_color=self.accent_color, hover_color=self.c("#2E3A6E"),
                      command=lambda: self.edit_note(n)).pack(side="left", padx=2)
        ctk.CTkButton(btn_frame, text="Xóa", width=50, fg_color="red",hover_color=self.c("#2E3A6E"), font=("Garet Variable", 13),
                      command=lambda: self.delete_note(n)).pack(side="left", padx=2)
        card.bind("<Button-1>", lambda e: self.edit_note(n))

    def upload_file(self):
        file_path = filedialog.askopenfilename(title="Chọn tài liệu", filetypes=[("Tài liệu", "*.pdf *.docx *.txt")])
        if not file_path:
            return
        
        if not os.path.exists(file_path):
            messagebox.showerror("Lỗi", "File không tồn tại!")
            return
        
        f_name = os.path.basename(file_path)
        if not os.path.exists(self.upload_dir):
            os.makedirs(self.upload_dir, exist_ok=True)
        
        import re
        safe_name = re.sub(r'[\\/*?:"<>|]', '', f_name)
        safe_name = safe_name.replace(" ", "_")[:50]
        dest_path = os.path.join(self.upload_dir, f"{int(time.time())}_{safe_name}")
        
        try:
            with open(file_path, 'rb') as source:
                with open(dest_path, 'wb') as target:
                    target.write(source.read())
        except Exception as e:
            messagebox.showerror("Lỗi", f"Không thể đọc file: {str(e)}")
            return
        
        size_kb = os.path.getsize(dest_path) // 1024
        size_str = f"{size_kb}KB" if size_kb < 1024 else f"{size_kb/1024:.1f}MB"
        ext = os.path.splitext(safe_name)[1].lower()
        loai = "pdf" if ext == ".pdf" else "doc"
        
        id_mon = None
        for ky, subjects in self.subjects_data.items():
            for sub in subjects:
                if sub['name'] == self.current_subject:
                    id_mon = sub['id']
                    break
            if id_mon:
                break
        
        if not id_mon:
            messagebox.showerror("Lỗi", "Không tìm thấy ID môn học!")
            return
        
        try:
            result = self.db.add_file(id_mon, safe_name, loai, dest_path)
            if result:
                self.load_data_from_db()
                messagebox.showinfo("Thành công", f"Đã tải lên {safe_name}")
                self.render_view()
            else:
                messagebox.showerror("Lỗi", "Upload thất bại!")
        except Exception as e:
            messagebox.showerror("Lỗi", f"Upload thất bại: {str(e)}")

    def upload_file_to_note(self):
        file_path = filedialog.askopenfilename(title="Chọn file để tạo ghi chú", filetypes=[("Văn bản", "*.txt *.docx")])
        if not file_path:
            return
        
        f_name = os.path.basename(file_path)
        import re
        safe_name = re.sub(r'[\\/*?:"<>|]', '', f_name)
        safe_name = safe_name.replace(" ", "_")
        ext = os.path.splitext(safe_name)[1].lower()
        
        content = ""
        try:
            if ext == ".txt":
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
            elif ext == ".docx" and HAS_LIBS:
                d = docx.Document(file_path)
                content = "\n".join([p.text for p in d.paragraphs])
            else:
                content = f"Nội dung từ file {safe_name}"
        except Exception as e:
            content = f"Lỗi đọc file: {str(e)}"
        
        try:
            self.db.add_note(self.current_subject, safe_name, content)
            self.load_data_from_db()
            messagebox.showinfo("Thành công", f"Đã tạo ghi chú từ file: {safe_name}")
            self.render_view()
        except Exception as e:
            messagebox.showerror("Lỗi", f"Tạo ghi chú thất bại: {e}")

    def create_new_note(self):
        self.current_note = {"_id": None, "title": "", "content": "", "time": "Vừa xong"}
        self.set_view_state("NOTE_EDITOR")

    def edit_note(self, note):
        self.current_note = note
        self.set_view_state("NOTE_EDITOR")

    def delete_note(self, note):
        if messagebox.askyesno("Xác nhận", f"Xóa ghi chú '{note['title']}'?"):
            try:
                self.db.delete_note(int(note['_id']))
                self.load_data_from_db()
                self.render_view()
            except Exception as e:
                messagebox.showerror("Lỗi", f"Xóa thất bại: {e}")

    def switch_to_viewer(self, file_info):
        self.current_file = file_info
        self.current_note = None
        self.set_view_state("VIEWER")

    # ==================== FILE VIEWER ====================
    def draw_inline_file_viewer(self):
        header = ctk.CTkFrame(self.main_view, fg_color=self.accent_color, height=65, corner_radius=0)
        header.pack(fill="x", side="top")

        ctk.CTkButton(header, text="← Quay lại", width=80,fg_color=self.accent_color, hover_color=self.c("#2E3A6E"), text_color="white",
                      font=("Garet Variable", 13), command=lambda: self.set_view_state("SUBJECT")).pack(side="left", padx=(20, 10))
        ctk.CTkLabel(header, text=f"📄 {self.current_file['name']}", text_color="white", 
                     font=("Garet Variable", 14, "bold")).pack(side="left")

        ctk.CTkButton(header, text="◧ Chia đôi màn hình", fg_color=self.c("#2E3A6E"), hover_color=self.c("#E8E8E8"), text_color="white",
                      font=("Garet Variable", 12, "bold"), corner_radius=15, height=38, width=180,
                      command=self.enter_split_mode).pack(side="right", padx=25)

        if self.current_file.get('path', '').lower().endswith('.pdf'):
            zoom_frame = ctk.CTkFrame(header, fg_color="transparent")
            zoom_frame.pack(side="right", padx=10)
            ctk.CTkButton(zoom_frame, text="+", width=30, height=30, fg_color=self.c("#2E3A6E"), hover_color=self.c("#E8E8E8"),
                          font=("Garet Variable", 13), command=self.zoom_in).pack(side="right", padx=2)
            ctk.CTkLabel(zoom_frame, text="Thu phóng", text_color="white", font=("Garet Variable", 12)).pack(side="right", padx=5)
            ctk.CTkButton(zoom_frame, text="-", width=30, height=30, fg_color=self.c("#2E3A6E"), hover_color=self.c("#E8E8E8"),
                          font=("Garet Variable", 13), command=self.zoom_out).pack(side="right", padx=2)

        scroll_v = ctk.CTkScrollableFrame(self.main_view, fg_color=self.c("#E8E8E8"), corner_radius=0)
        scroll_v.pack(fill="both", expand=True)
        self.render_file_content(scroll_v, self.current_file)

    def enter_split_mode(self):
        self.user_set_split_pos = None
        self.current_columns = 4
        self.set_view_state("SPLIT")

    def zoom_in(self):
        self.pdf_zoom += 0.2
        self.render_view()

    def zoom_out(self):
        if self.pdf_zoom > 0.4:
            self.pdf_zoom -= 0.2
            self.render_view()

    def render_file_content(self, parent, f_info):
        for w in parent.winfo_children():
            w.destroy()
        path = f_info.get('path', '')
        if not path or not os.path.exists(path):
            ctk.CTkLabel(parent, text="⚠️ File không tồn tại", font=("Garet Variable", 16, "bold"), text_color="#E74C3C").pack(pady=200)
            return
        if not HAS_LIBS:
            ctk.CTkLabel(parent, text="⚠️ Thiếu thư viện đọc file", font=("Garet Variable", 16), text_color="red").pack(pady=200)
            return
        ext = os.path.splitext(path)[1].lower()
        try:
            if ext == ".pdf":
                doc = fitz.open(path)
                for i in range(min(30, len(doc))):
                    page = doc.load_page(i)
                    pix = page.get_pixmap(matrix=fitz.Matrix(self.pdf_zoom, self.pdf_zoom))
                    img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                    ctk_img = ctk.CTkImage(light_image=img, size=(pix.width, pix.height))
                    lbl = ctk.CTkLabel(parent, image=ctk_img, text="")
                    lbl.pack(pady=10)
                doc.close()
            elif ext == ".docx":
                d = docx.Document(path)
                full_text = "\n".join([p.text for p in d.paragraphs])
                txt = ctk.CTkTextbox(parent, font=("Garet Variable", 14), fg_color="white", text_color="black")
                txt.pack(fill="both", expand=True, padx=20, pady=20)
                txt.insert("0.0", full_text)
                txt.configure(state="disabled")
            elif ext == ".txt":
                with open(path, 'r', encoding='utf-8') as f:
                    full_text = f.read()
                txt = ctk.CTkTextbox(parent, font=("Garet Variable", 14), fg_color="white", text_color="black")
                txt.pack(fill="both", expand=True, padx=20, pady=20)
                txt.insert("0.0", full_text)
                txt.configure(state="disabled")
        except Exception as e:
            ctk.CTkLabel(parent, text=f"Lỗi: {str(e)}", text_color="red").pack(pady=100)

    # ==================== NOTE EDITOR ====================
    def draw_inline_note_editor(self):
        header = ctk.CTkFrame(self.main_view, fg_color=self.accent_color, height=65, corner_radius=0)
        header.pack(fill="x", side="top")
        ctk.CTkButton(header, text="← Quay lại", width=80, fg_color=self.accent_color, hover_color=self.c("#2E3A6E"), text_color="white",
                      font=("Garet Variable", 13), command=lambda: [setattr(self, 'current_note', None), self.set_view_state("SUBJECT")]).pack(side="left", padx=(20, 10))
        ctk.CTkLabel(header, text="📝 Soạn thảo Ghi chú", text_color="white", font=("Garet Variable", 13, "bold")).pack(side="left", padx=10)
        editor_frame = ctk.CTkFrame(self.main_view, fg_color=self.bg_color, corner_radius=0)
        editor_frame.pack(fill="both", expand=True)
        self.draw_note_editor_ui(editor_frame, self.current_note)

    def draw_note_editor_ui(self, parent, note_info):
        title_entry = ctk.CTkEntry(parent, placeholder_text="Nhập tiêu đề ghi chú...", height=45,
                                   font=("Garet Variable", 18, "bold"), fg_color="white", border_width=0)
        title_entry.pack(fill="x", padx=20, pady=(20, 10))
        if note_info and note_info.get('title'):
            title_entry.insert(0, note_info['title'])
        note_body = ctk.CTkTextbox(parent, font=("Garet Variable", 15), corner_radius=10, 
                                   border_width=1, border_color="#E0E0E0")
        note_body.pack(fill="both", expand=True, padx=20, pady=(0, 20))
        if note_info and note_info.get('content'):
            note_body.insert("0.0", note_info['content'])
        btn_frame = ctk.CTkFrame(parent, fg_color="transparent")
        btn_frame.pack(pady=20)
        ctk.CTkButton(btn_frame, text="💾 Lưu ghi chú", fg_color="#27AE60", height=40, width=150,
                      font=("Garet Variable", 13, "bold"),
                      command=lambda: self.save_note(title_entry.get(), note_body.get("1.0", "end-1c"))).pack(side="left", padx=10)
        ctk.CTkButton(btn_frame, text="❌ Hủy", fg_color="gray", height=40, width=100,
                      font=("Garet Variable", 13), command=lambda: self.set_view_state("SUBJECT")).pack(side="left", padx=10)

    def save_note(self, title, content):
        if not title.strip():
            messagebox.showwarning("Thông báo", "Vui lòng nhập tiêu đề ghi chú!")
            return
        if self.current_note and self.current_note.get('_id'):
            try:
                self.db.update_note(int(self.current_note['_id']), title, content)
                for n in self.notes_data.get(self.current_subject, []):
                    if n['_id'] == self.current_note['_id']:
                        n['title'] = title
                        n['content'] = content
                        break
                messagebox.showinfo("Thành công", f"Đã cập nhật ghi chú '{title}'")
            except Exception as e:
                messagebox.showerror("Lỗi", f"Cập nhật thất bại: {e}")
        else:
            try:
                new_id = self.db.add_note(self.current_subject, title, content)
                if new_id:
                    if self.current_subject not in self.notes_data:
                        self.notes_data[self.current_subject] = []
                    self.notes_data[self.current_subject].append({
                        "_id": new_id,
                        "title": title,
                        "content": content,
                        "time": datetime.now().strftime("%d/%m/%Y")
                    })
                    messagebox.showinfo("Thành công", f"Đã tạo ghi chú '{title}'")
                else:
                    messagebox.showerror("Lỗi", "Tạo ghi chú thất bại!")
                    return
            except Exception as e:
                messagebox.showerror("Lỗi", f"Tạo ghi chú thất bại: {e}")
                return
        self.current_note = None
        self.set_view_state("SUBJECT")

    # ==================== SPLIT VIEW ====================
    def draw_inline_split_session(self):
        toolbar = ctk.CTkFrame(self.main_view, fg_color="#323639", height=50, corner_radius=0)
        toolbar.pack(fill="x")
        ctk.CTkButton(toolbar, text="← Thoát chế độ chia đôi", fg_color=self.c("#2E3A6E"), hover_color=self.c("#E8E8E8"), text_color="white",
                      width=180, font=("Garet Variable", 10, "bold"), command=self.exit_split_mode).pack(side="left", padx=(20, 10))

        self.paned_window = PanedWindow(self.main_view, orient="horizontal", bg="#DDDDDD", 
                                         bd=0, sashwidth=6, sashrelief="flat")
        self.paned_window.pack(fill="both", expand=True)

        left_side = ctk.CTkFrame(self.paned_window, fg_color="#525659", corner_radius=0)
        self.paned_window.add(left_side)

        if self.current_file:
            self.draw_split_left_with_file(left_side)
        else:
            self.draw_split_left_file_list(left_side)

        right_side = ctk.CTkFrame(self.paned_window, fg_color=self.bg_color, corner_radius=0)
        self.paned_window.add(right_side)

        if self.current_note is not None:
            self.draw_note_editor_ui(right_side, self.current_note)
        else:
            self.draw_split_right_note_list(right_side)

        def on_sash_moved(e):
            if hasattr(self, 'paned_window') and self.paned_window.winfo_exists():
                self.user_set_split_pos = self.paned_window.sash_coord(0)[0]
        self.paned_window.bind("<<PanedWindow::SashMoved>>", on_sash_moved)

        if self.user_set_split_pos:
            self.after(100, self.restore_split_pos)
        else:
            self.after(100, self.adjust_split_50_50)

    def adjust_split_50_50(self):
        if hasattr(self, 'paned_window') and self.paned_window.winfo_exists():
            self.paned_window.update_idletasks()
            total_w = self.paned_window.winfo_width()
            if total_w > 100:
                self.paned_window.sash_place(0, total_w // 2, 0)

    def restore_split_pos(self):
        if hasattr(self, 'paned_window') and self.paned_window.winfo_exists() and self.user_set_split_pos:
            total_w = self.paned_window.winfo_width()
            if total_w > 100 and self.user_set_split_pos < total_w:
                self.paned_window.sash_place(0, self.user_set_split_pos, 0)

    def exit_split_mode(self):
        self.current_file = None
        self.current_note = None
        self.user_set_split_pos = None
        if hasattr(self, 'paned_window'):
            self.paned_window.destroy()
        self.set_view_state("SUBJECT")

    def draw_split_left_with_file(self, parent):
        bar = ctk.CTkFrame(parent, fg_color="#2B2B2B", height=40, corner_radius=0)
        bar.pack(fill="x")
        ctk.CTkButton(bar, text="✖ Đóng", width=60, fg_color="transparent", text_color="white", 
                      hover_color="#E74C3C", font=("Garet Variable", 13), command=self.close_current_file).pack(side="left", padx=5)
        ctk.CTkLabel(bar, text=self.current_file['name'], text_color="white", font=("Garet Variable", 12)).pack(side="left", padx=5)

        if self.current_file.get('path', '').lower().endswith('.pdf'):
            zoom_frame = ctk.CTkFrame(bar, fg_color="transparent")
            zoom_frame.pack(side="right", padx=10)
            ctk.CTkButton(zoom_frame, text="+", width=30, height=30, fg_color="#555", hover_color="#777",
                          font=("Garet Variable", 13), command=self.zoom_in_split).pack(side="right", padx=2)
            ctk.CTkButton(zoom_frame, text="-", width=30, height=30, fg_color="#555", hover_color="#777",
                          font=("Garet Variable", 13), command=self.zoom_out_split).pack(side="right", padx=2)

        viewer_scroll = ctk.CTkScrollableFrame(parent, fg_color="#E8E8E8", corner_radius=0)
        viewer_scroll.pack(fill="both", expand=True)
        self.render_file_content(viewer_scroll, self.current_file)

    def close_current_file(self):
        self.current_file = None
        self.render_view()

    def zoom_in_split(self):
        self.pdf_zoom += 0.2
        self.render_view()

    def zoom_out_split(self):
        if self.pdf_zoom > 0.4:
            self.pdf_zoom -= 0.2
            self.render_view()

    def draw_split_left_file_list(self, parent):
        top_bar = ctk.CTkFrame(parent, fg_color="#F0F0F0", corner_radius=0)
        top_bar.pack(fill="x", pady=0)
        ctk.CTkLabel(top_bar, text="📂 DANH SÁCH TÀI LIỆU", font=("Garet Variable", 16, "bold"), 
                     text_color="black").pack(side="left", padx=20, pady=15)
        ctk.CTkButton(top_bar, text="📥 Upload", fg_color=self.c("#2E3A6E"), hover_color=self.c("#E8E8E8"), height=32, width=80,
                      command=self.upload_file).pack(side="right", padx=20)

        scroll = ctk.CTkScrollableFrame(parent, fg_color="#E8E8E8", corner_radius=0)
        scroll.pack(fill="both", expand=True)

        files = self.files_data.get(self.current_subject, [])
        for f in files:
            btn = ctk.CTkButton(scroll, text="📄 " + f['name'], fg_color="white", text_color="black", 
                                height=55, anchor="w", hover_color=self.c("#D1D5F0"), font=("Garet Variable", 13, "bold"),
                                command=lambda file=f: self.open_file_in_split(file))
            btn.pack(fill="x", pady=5, padx=10)

    def open_file_in_split(self, file):
        self.current_file = file
        self.render_view()

    def draw_split_right_note_list(self, parent):
        top_bar = ctk.CTkFrame(parent, fg_color="transparent")
        top_bar.pack(fill="x", padx=20, pady=(20, 10))
        ctk.CTkLabel(top_bar, text="📝 DANH SÁCH GHI CHÚ", font=("Garet Variable", 18, "bold"), 
                     text_color=self.accent_color).pack(side="left")
        ctk.CTkButton(top_bar, text="+ Thêm", fg_color=self.c("#2E3A6E"), hover_color=self.c("#E8E8E8"), height=35, width=90,
                      font=("Garet Variable", 13, "bold"), command=self.create_new_note_split).pack(side="right")
        ctk.CTkButton(top_bar, text="📤 Upload file", fg_color=self.c("#2E3A6E"), hover_color=self.c("#E8E8E8"), height=35, width=100, text_color="white",
                      font=("Garet Variable", 13, "bold"), command=self.upload_file_to_note).pack(side="right", padx=10)

        scroll = ctk.CTkScrollableFrame(parent, fg_color="transparent")
        scroll.pack(fill="both", expand=True, padx=20, pady=10)

        notes = self.notes_data.get(self.current_subject, [])
        for n in notes:
            btn_frame = ctk.CTkFrame(scroll, fg_color="transparent")
            btn_frame.pack(fill="x", pady=5)
            btn = ctk.CTkButton(btn_frame, text="📝 " + n['title'], fg_color="white", text_color="black",
                                height=55, anchor="w", hover_color=self.c("#D1D5F0"), font=("Garet Variable", 13, "bold"),
                                command=lambda note=n: self.open_note_in_split(note))
            btn.pack(side="left", fill="x", expand=True)
            ctk.CTkButton(btn_frame, text="✎", width=40, fg_color="#FF9800",
                          font=("Garet Variable", 13), command=lambda note=n: self.open_note_in_split(note)).pack(side="right", padx=2)

    def create_new_note_split(self):
        self.current_note = {"_id": None, "title": "", "content": "", "time": "Vừa xong"}
        self.render_view()

    def open_note_in_split(self, note):
        self.current_note = note
        self.render_view()
        
    def adjust_layout(self):
        if self.view_state == "SPLIT" and hasattr(self, 'paned_window') and self.paned_window.winfo_exists():
            self.after(100, self.adjust_split_50_50)
        elif self.view_state == "HOME":
            # Cập nhật lại giao diện khi thay đổi kích thước
            self.render_view()
        else:
            self.render_view()
